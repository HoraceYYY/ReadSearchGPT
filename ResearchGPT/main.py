from Google import async_google
from aiohttp import ClientSession
from docx import Document
from bs4 import BeautifulSoup
import time, os
from fastapi import FastAPI, BackgroundTasks, Depends, HTTPException, Cookie, status
from fastapi.responses import FileResponse, Response
from pydantic import BaseModel
from typing import List
from fastapi.middleware.cors import CORSMiddleware # to allow CORS
from database import SessionLocal
from sqlalchemy.orm import Session
from sqlalchemy import desc
import models, secrets, logging

class EmailRequest(BaseModel):
    research_id: str
    emails: List[str]

class APIKey(BaseModel):
    apiKey: str

class ResearchID(BaseModel):
    taskID: str

class additionalSearch(BaseModel):
    queryID: str
    apiKey: str

class firstSearch(BaseModel):
    searchqueries: List[str]
    searchDomain: str | None = None
    apiKey: str

class deepsearch(BaseModel):
    queryID: str
    searchDomain: str | None = None
    apiKey: str

class DownloadResult(BaseModel):
    queryIDs: List[str]

class FeedbackBase(BaseModel):
    feedback: str

class UserResponse(BaseModel):
    userid: str | None = None

app = FastAPI()

# origins = [
#     "http://localhost:5173",  
#     "https://readsearch.azurewebsites.net",
#     "www.readsearchgpt.com",
#     "https://readsearchgpt.com"
# ]

# app.add_middleware(
#     CORSMiddleware,
#     allow_origins=origins,
#     allow_credentials=True,
#     allow_methods=["*"],
#     allow_headers=["*"],
# )
# Get a db session
def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()

@app.post("/firstsearch") # this is the entry point of the search, this will search all the topics entered
async def first_search(search: firstSearch, userid: str = Cookie(None), db: Session = Depends(get_db)):
    logging.info(f"Userid: {userid}")
    try:
        if not userid:
            userid = "user id not logged"
            raise HTTPException(status_code=status.HTTP_401_UNAUTHORIZED, detail="Userid cookie is missing")
        searchqueries = search.searchqueries
        userDomain = search.searchDomain
        api_key = search.apiKey
        querywithresult = await async_google.first_search(db, searchqueries, userDomain, api_key, userid)
        db.commit()
        print(f"Task Completed")
    finally:    
        db.close()
    return querywithresult, api_key# returning api key to store in the front end to make the second and deep search, returning the userDomain for the deep search

@app.post("/secondsearch")
async def second_search(search: additionalSearch, db: Session = Depends(get_db)):
    try:
        queryid = search.queryID
        api_key = search.apiKey
        querywithresult = await async_google.second_search(db, queryid, api_key) # this results includes the initial search as well
        db.commit()
        print(f"Task Completed")
    finally:
        db.close()
    return querywithresult, api_key # returning api key to store in the front end to make the next search

@app.post("/firstdeepsearch/")
async def first_deep_search(search: deepsearch, db: Session = Depends(get_db)):
    try:
        queryid = search.queryID
        userDomain = search.searchDomain
        api_key = search.apiKey
        querywithresult = await async_google.first_deep_search(db, queryid, userDomain, api_key) # this results includes the initial search as well
        db.commit()
        print(f"Task Completed")
    finally:
        db.close()
    return querywithresult, api_key

async def download_word(queryids, db: Session = Depends(get_db)):

    document = Document() # Initialize a new Word document
    data = {}  # Initialize a dictionary to store the data

    if not queryids:
        raise HTTPException(status_code=404, detail="No queries found for this task ID")

    for query_id in queryids: # Fetch the Query, URLData and URLSummary from the database
        query = db.query(models.Query).filter(models.Query.id == query_id).first()
        url_data = db.query(models.URLData).filter(models.URLData.query_id == query_id).all()
        url_summary = db.query(models.URLSummary).filter(models.URLSummary.query_id == query_id).all()
        
        data[query_id] = { # Save the data in the dictionary
            'query': query.query,
            'url_data': [{'title': item.title, 'content': item.content, 'category': item.category, 'url': item.url} for item in url_data],
            'url_summary': [{'summarytype': item.summarytype, 'summary': item.summary} for item in url_summary]
        }
    
    document.add_heading('Research Summry', level=1)
    for query_id, info in data.items(): # Loop over the data dictionary to populate the Word document
        document.add_heading(info['query'].capitalize(), level=2) # Add the Query to the Word document
        for summary in info['url_summary']:
            document.add_heading(f'{summary["summarytype"]}', level=3)
            soup = BeautifulSoup(summary["summary"], 'html.parser')
            text = soup.get_text()
            document.add_paragraph(f'Summary: {text}')
    
    document.add_heading('Individual Web Results', level=1)
    for query_id, info in data.items():
        document.add_heading(info['query'].capitalize(), level=2)    
        # Add the URLData to the Word document
        for data in [d for d in info['url_data'] if d["category"] == 'Website_Content']:
            para = document.add_paragraph()
            run = para.add_run(f'{data["title"]}')
            run.bold = True
            document.add_paragraph(f'{data["url"]}')
            soup = BeautifulSoup(data["content"])
            text = soup.get_text()
            document.add_paragraph(f'{text}')
        para = document.add_paragraph()
        run = para.add_run('Additional PDFs')
        run.bold = True
        for data in [d for d in info['url_data'] if d["category"] == 'PDFs']:
            document.add_paragraph(f'URL: {data["url"]}')
        para = document.add_paragraph()
        run = para.add_run('Additional Relevant Websites')
        run.bold = True
        for data in [d for d in info['url_data'] if d["category"] == "Unread_Websites"]:
            document.add_paragraph(f'URL: {data["url"]}')
    # Save the document
    file_path = f"Results/Readsearch_Report.docx"
    document.save(file_path)
    return file_path

@app.post("/webdownload")
async def web_download_excel(download: DownloadResult, background_tasks: BackgroundTasks, db: Session = Depends(get_db)):
    queryids = download.queryIDs
    print(queryids)
    file_path = await download_word(queryids, db)
    background_tasks.add_task(delete_file_after_delay, file_path, delay=30) # Add a background task to delete the file after 1 minute
    return FileResponse(file_path, media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document', filename=f"Readsearch_Report.docx")  # Return the Word document as a response

def delete_file_after_delay(file_path: str, delay: int):
    time.sleep(delay)
    if os.path.exists(file_path):
        os.remove(file_path)

@app.post("/testapi")
async def testAPI(apiKey: APIKey):
    api_key = apiKey.apiKey
    url = 'https://api.openai.com/v1/chat/completions'
    headers = {
        'Content-Type': 'application/json',
        'Authorization': f'Bearer {api_key}'
    }
    body = {
        "model": "gpt-3.5-turbo",
        "messages": [{"role": "user", "content": "Hello!"}]
    }

    async with ClientSession() as session:
        try:
            async with session.post(url, headers=headers, json=body) as response:
                status_code = response.status
                if status_code == 200:
                    print("success: ", status_code)
                    return {"Key": "Valid"}
                else:
                    print("error: ", status_code)
                    return {"Key": f"Not Valid: {status_code}"}
        except Exception as e:
            print(e)
            return {"Key": f"Not Valid: {e}"}

@app.post("/feedback")
async def create_feedback(feedback: FeedbackBase, db: Session = Depends(get_db)):
    new_feedback = models.Feedback(
        feedback=feedback.feedback
    )
    db.add(new_feedback)
    db.commit()
    db.refresh(new_feedback)
    return {"Message": "Feedback successfully submitted", "Feedback ID": new_feedback.id}

@app.get("/create-cookie")
def create_cookie():
    response = Response()
    unique_id = secrets.token_urlsafe(16) # generates a unique identifier
    response.set_cookie(key="userid", value=unique_id, max_age=10*365*24*60*60, secure=True, samesite='Lax', httponly=True, domain='.readsearchgpt.com')
    #response.set_cookie(key="userid", value=unique_id)
    print("created cookie after check: ", unique_id)
    return response

@app.get("/read-cookie", response_model=UserResponse)
def read_cookie(userid: str | None = Cookie(None)):
    print("checking cookie user id", userid)
    return {"userid": userid}

@app.get("/queryhistory")
async def get_queries(userid: str = Cookie(None), db: Session = Depends(get_db)):
    if not userid:
        return None
    print(userid)
    tasks = db.query(models.Task).filter(models.Task.userid == userid).all()
    if not tasks:
        return None
    query_dict = {}
    for task in tasks:
        queries = db.query(models.Query).filter(models.Query.task_id == task.id).all()
        for query in queries:
            query_dict[str(query.id)] = query.query

    return query_dict

@app.post("/historicalresults")
async def get_historical_results(queryids: DownloadResult, db: Session = Depends(get_db)):
    query_ids = queryids.queryIDs
    print(query_ids)
    queryresults = {}
    for query_id in query_ids:
        query_object = db.query(models.Query.query).filter(models.Query.id == query_id).first()
        url_data_objects = db.query(models.URLData.content, models.URLData.title, models.URLData.url).filter(models.URLData.query_id == query_id).all()
        url_summary_object = db.query(models.URLSummary.summary).filter(models.URLSummary.query_id == query_id).order_by(desc(models.URLSummary.timestamp)).first()

        if query_object and url_data_objects and url_summary_object:
            query = query_object.query
            results = [{"content": obj[0], "title": obj[1], "url": obj[2]} for obj in url_data_objects]
            summary = url_summary_object.summary
            results.append({"Summary": summary})
            queryresults[query_id] = [query, results]

    return queryresults, "" ## "" represents an empty api key